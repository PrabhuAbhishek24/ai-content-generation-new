from flask import Flask, request, jsonify, send_file
import openai
import os
import zipfile
import io
from fpdf import FPDF
from docx import Document
from pathlib import Path
from dotenv import load_dotenv
from docx.shared import Inches

app = Flask(__name__)

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")


# Function to fetch response from GPT
def fetch_gpt_response(query):
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in the pharmaceutical and medical domain only. Only answer those questions and don't answer any other questions."},
                {"role": "user", "content": query},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"


def save_as_scorm_pdf(content, output_folder="scorm_package", scorm_zip_name="scorm_package.zip"):
    # Step 1: Create the SCORM folder structure
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Save the PDF
    pdf_file_path = os.path.join(output_folder, "content.pdf")
    save_as_pdf(content, pdf_file_path)

    # Step 2: Create the HTML file
    html_file_path = os.path.join(output_folder, "index.html")
    with open(html_file_path, "w", encoding="utf-8") as html_file:
        html_file.write(f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>SCORM Content</title>
        </head>
        <body>
            <h1>Research Content Response</h1>
            <iframe src="content.pdf" width="100%" height="600px"></iframe>
        </body>
        </html>
        """)

    # Step 3: Create the imsmanifest.xml file
    manifest_file_path = os.path.join(output_folder, "imsmanifest.xml")
    with open(manifest_file_path, "w", encoding="utf-8") as manifest_file:
        manifest_file.write(f"""
        <?xml version="1.0" encoding="UTF-8"?>
        <manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
                  xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_v1p3"
                  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
                  xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1
                                      http://www.imsglobal.org/xsd/imscp_v1p1.xsd
                                      http://www.adlnet.org/xsd/adlcp_v1p3
                                      http://www.adlnet.org/xsd/adlcp_v1p3.xsd">
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <organizations>
                <organization identifier="ORG-1">
                    <title>Research Content</title>
                    <item identifier="ITEM-1" identifierref="RES-1">
                        <title>Research Content Response</title>
                    </item>
                </organization>
            </organizations>
            <resources>
                <resource identifier="RES-1" type="webcontent" href="index.html">
                    <file href="index.html"/>
                    <file href="content.pdf"/>
                </resource>
            </resources>
        </manifest>
        """)

    # Step 4: Zip the SCORM package
    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        for foldername, subfolders, filenames in os.walk(output_folder):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, output_folder)
                scorm_zip.write(file_path, arcname)

    # Provide the download button for the SCORM package
    with open(scorm_zip_name, 'rb') as scorm_file:
        return scorm_file.read()


def save_as_pdf(content, file_name="response.pdf"):
    pdf = FPDF()
    pdf.add_page()

    # Add the logo
    pdf.image("assets/logo.jpeg", x=10, y=8, w=30)

    # Title of the document
    pdf.set_font("Arial", style='B', size=16)
    pdf.ln(30)
    pdf.cell(200, 10, txt="Research Content Response", ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, content)

    # Save the PDF
    pdf.output(file_name)


def save_as_scorm_word(content, file_name="scorm_package.zip"):
    # Create an in-memory zip file
    scorm_zip = io.BytesIO()

    with zipfile.ZipFile(scorm_zip, 'w') as zf:
        # Create and add manifest.xml
        manifest_content = """<manifest>
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <resources>
                <resource identifier="res1" type="webcontent" href="response.docx">
                    <file href="response.docx"/>
                    <file href="response.html"/>
                </resource>
            </resources>
        </manifest>"""
        zf.writestr("imanifest.xml", manifest_content)

        # Create DOCX file
        docx_buffer = io.BytesIO()
        doc = Document()
        # Add the logo to the Word document
        logo_path = "assets/logo.jpeg"
        if Path(logo_path).is_file():
            doc.add_picture(logo_path, width=Inches(1.5))
        doc.add_paragraph('\n')
        doc.add_paragraph("Research Content Response", style='Heading 1')
        doc.add_paragraph('\n')
        doc.add_paragraph(content)
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        zf.writestr("response.docx", docx_buffer.getvalue())

        # Create HTML file
        html_content = f"""
        <html>
        <head><title>Research Content Response</title></head>
        <body>
        <h1>Research Content Response</h1>
        <p>{content.replace('\n', '<br>')}</p>
        </body>
        </html>
        """
        zf.writestr("index.html", html_content)

    scorm_zip.seek(0)
    return scorm_zip.getvalue()





# API route for generating content
@app.route('/api/generate-content', methods=['POST'])
def generate_content():
    try:
        # Get the query from the request JSON
        data = request.get_json()
        query = data.get("query", "").strip()

        if not query:
            return jsonify({"error": "Query is required."}), 400

        # Process the query for content generation
        gpt_response = fetch_gpt_response(query)

        return jsonify({"query": query, "response": gpt_response})

    except Exception as e:
        return jsonify({"error": f"Internal server error: {str(e)}"}), 500


# API route for downloading content in different formats
@app.route('/api/download-content', methods=['POST'])
def download_content():
    try:
        # Get the data from the request JSON
        data = request.get_json()
        response = data.get("response")
        format_type = data.get("format")

        if not response or not format_type:
            return jsonify({"error": "Response and format are required."}), 400

        if format_type == "pdf_scorm":
            scorm_file = save_as_scorm_pdf(response, "pdf")
            return send_file(io.BytesIO(scorm_file), as_attachment=True, download_name="scorm_package_pdf.zip")
        elif format_type == "docx_scorm":
            scorm_file = save_as_scorm_word(response, "docx")
            return send_file(io.BytesIO(scorm_file), as_attachment=True, download_name="scorm_package_doc.zip")
        else:
            return jsonify({"error": "Invalid download format selected."}), 400

    except Exception as e:
        return jsonify({"error": f"Error in downloading content: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True)
